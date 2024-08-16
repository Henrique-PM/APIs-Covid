from flask import Flask, render_template, request, jsonify, send_file
import os
import sqlite3
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests

app = Flask(__name__)

rest_api = "https://restcountries.com/v3.1/name/"
covid_api = "https://disease.sh/v3/covid-19/countries/"
vaccine_api = "https://disease.sh/v3/covid-19/vaccine/coverage/countries/"
morte_api = "https://disease.sh/v3/covid-19/historical/all"
noticia_api = "https://newsapi.org/v2/everything?q=COVID-19&apiKey=8d5c096f74b440db886fc179e3de06af"

def insert_health_data(pais, populacao, casos, mortes, pessoas_vacinadas):
    try:
        with sqlite3.connect('bancoInfo.db') as conn:
            c = conn.cursor()
            c.execute('''
            INSERT INTO info (pais, populacao, casos, pessoas_vacinadas, mortes)
            VALUES (?, ?, ?, ?, ?)
            ''', (pais, populacao, casos, pessoas_vacinadas, mortes))
            conn.commit()
            return {"status": "success", "message": "Data inserted successfully"}
    except sqlite3.Error as e:
        return {"status": "error", "message": str(e)}
    
@app.route('/save-health-data', methods=['POST'])
def save_health_data():
    data = request.get_json()
    pais = data.get('pais')
    populacao = data.get('populacao')
    casos = data.get('casos')
    mortes = data.get('mortes')
    pessoas_vacinadas = data.get('pessoas_vacinadas')
    result = insert_health_data(pais, populacao, casos, mortes, pessoas_vacinadas)
    return jsonify(result)

@app.route("/")
def index():
    api_noticia = requests.get(noticia_api)

    if api_noticia.status_code == 200:
        noticia_data = api_noticia.json()
        articles = noticia_data['articles'][:12]
        return render_template('index.html', articles=articles)
    return jsonify({"error": "Erro ao obter notícias"}), 500

@app.route("/dados_saude", methods=["POST"])
def get_health_data():
    pais = request.form.get("country_name")

    if not pais:
        return jsonify({"error": "Nome do país não fornecido"}), 400

    try:
        rest_link = f"{rest_api}{pais}"
        api_rest = requests.get(rest_link)
        covid_link = f"{covid_api}{pais}"
        api_covid = requests.get(covid_link)
        vaccine_link = f"{vaccine_api}{pais}?lastdays=1&fullData=false"
        api_vaccine = requests.get(vaccine_link)
        morte_link = f"{morte_api}{pais}"
        api_morte = requests.get(morte_link)

        population, casos, pessoas_vacinadas, deaths = None, None, None, None

        if api_rest.status_code == 200:
            population_rest = api_rest.json()
            population = population_rest[0].get("population", "Não disponível")

        if api_covid.status_code == 200:
            casos_covid = api_covid.json()
            casos = casos_covid.get("cases", "Não disponível")

        if api_vaccine.status_code == 200:
            vaccine_data = api_vaccine.json()
            pessoas_vacinadas = vaccine_data["timeline"].get(
                str(list(vaccine_data["timeline"].keys())[0]), "Não disponível")

        if api_morte.status_code == 200:
            casos_morte = api_morte.json()
            deaths = casos_morte.get("deaths", "Não disponível")

        if not (population and casos and pessoas_vacinadas and deaths):
            return jsonify({"error": "Erro ao obter todos os dados necessários"}), 500

        # Debug: Print the values to ensure they are correct
        print(f"Dados para salvar: {pais}, {population}, {casos}, {pessoas_vacinadas}, {deaths}")

        return jsonify({"message": "Dados salvos com sucesso"}), 200

    except Exception as e:
        return jsonify({"error": f"Erro ao processar dados: {str(e)}"}), 500

def get_document(doc_id):
    document_data = {
        'pais': '', 'populacao': '', 'casos': '', 'pessoas_vacinadas': '',
        'mortes': ''
    }
    try:
        with sqlite3.connect('bancoInfo.db') as conn:
            c = conn.cursor()
            c.execute("SELECT pais,populacao,casos,pessoas_vacinadas,mortes FROM info WHERE id=?", (doc_id,))
            result = c.fetchone()
            if result:
                document_data['pais'], document_data['populacao'], document_data['casos'], document_data['pessoas_vacinadas'], document_data['mortes'] = result
    except sqlite3.Error as e:
        print(f"Erro ao obter dados do banco de dados: {e}")
    return document_data

def formatar_documento_abnt(doc_id):
    document_data = get_document(doc_id)

    if not any(document_data.values()):
        return None

    doc = Document()
    estilo_paragrafo = doc.styles['Normal']
    estilo_paragrafo.font.name = 'Times New Roman'
    estilo_paragrafo.font.size = Pt(12)
    estilo_paragrafo.paragraph_format.space_before = Pt(6)
    estilo_paragrafo.paragraph_format.space_after = Pt(6)
    estilo_paragrafo.paragraph_format.line_spacing = 1.5
    estilo_paragrafo.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    for key, value in document_data.items():
        if value:
            doc.add_paragraph(f"{key.capitalize()}:")
            doc.add_paragraph(value, style='Normal')
            doc.add_paragraph()

    secao = doc.sections[0]
    secao.left_margin = Pt(72)
    secao.right_margin = Pt(72)
    secao.top_margin = Pt(72)
    secao.bottom_margin = Pt(72)

    static_dir = os.path.join(os.getcwd(), 'static')
    if not os.path.exists(static_dir):
        os.makedirs(static_dir)

    nome_arquivo = os.path.join(static_dir, f"documento_{doc_id}.docx")
    doc.save(nome_arquivo)

    return nome_arquivo

@app.route('/baixar', methods=['POST'])
def baixar():
    doc_id = request.form.get('id')

    if not doc_id:
        return jsonify({"error": "ID do documento não fornecido"}), 400

    documento_salvo = formatar_documento_abnt(doc_id)
    if documento_salvo and os.path.exists(documento_salvo):
        return send_file(documento_salvo, as_attachment=True)
    
    return jsonify({"error": "Erro ao baixar documento"}), 500

if __name__ == '__main__':
    app.run(debug=True)
